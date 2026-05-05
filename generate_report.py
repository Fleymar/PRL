from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Helpers ───────────────────────────────────────────────────────────────────
BLUE  = RGBColor(0x1F, 0x49, 0x7D)
GREY  = RGBColor(0x40, 0x40, 0x40)
LGREY = RGBColor(0x70, 0x70, 0x70)

def set_font(run, name="Calibri", size=11, bold=False, color=None, italic=False):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, size=15, bold=True, color=BLUE)
    # underline via border on paragraph bottom
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=12, bold=True, color=GREY)
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    set_font(run, size=11, color=GREY)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.6)
    run = p.add_run(text)
    set_font(run, size=11, color=GREY)
    return p

def kv(key, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.8)
    r1 = p.add_run(key + " : ")
    set_font(r1, size=11, bold=True, color=GREY)
    r2 = p.add_run(value)
    set_font(r2, size=11, color=LGREY)
    return p

def separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_comparison_table(rows):
    table = doc.add_table(rows=len(rows)+1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["Composante", "Version 1 (v1)", "Version 2 (v2)"]):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        set_font(run, size=10, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F497D")
        tcPr.append(shd)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1].cells
        for ci, cell_text in enumerate(row_data):
            row[ci].text = cell_text
            run = row[ci].paragraphs[0].runs[0]
            bold = ci == 0
            set_font(run, size=10, bold=bold, color=GREY)
            if ri % 2 == 0:
                tc = row[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F2F2F2")
                tcPr.append(shd)
    return table


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE DE GARDE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Rapport de Jalon 3")
set_font(r, size=26, bold=True, color=BLUE)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("Projet PRL — Agent Rocket League par Apprentissage par Renforcement")
set_font(r, size=13, bold=False, color=GREY, italic=True)

doc.add_paragraph()
separator()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(f"Date : {datetime.date.today().strftime('%d %B %Y')}   |   Version : v2.0")
set_font(r, size=11, color=LGREY)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. CONTEXTE
# ══════════════════════════════════════════════════════════════════════════════
heading1("1. Contexte et objectifs du jalon 3")
separator()

body(
    "Ce rapport documente les travaux réalisés dans le cadre du troisième jalon du projet PRL. "
    "L'objectif de ce jalon était de dépasser les limitations identifiées lors du jalon 2 en procédant "
    "à une refonte complète de l'architecture d'entraînement, fondée sur les meilleures pratiques "
    "identifiées au sein de la communauté rlgym et de la littérature en apprentissage par renforcement "
    "appliqué aux jeux vidéo."
)
body(
    "Les axes d'amélioration ciblés sont : l'espace d'actions, la fonction de récompense, "
    "la stratégie adversariale et l'architecture réseau. Chacun de ces axes a fait l'objet "
    "d'une recherche documentaire approfondie avant implémentation."
)

# ══════════════════════════════════════════════════════════════════════════════
#  2. BILAN V1
# ══════════════════════════════════════════════════════════════════════════════
heading1("2. Bilan de la version 1 (pré-v2)")
separator()

heading2("2.1 Architecture initiale")
body(
    "La version 1 de l'agent reposait sur une architecture simple, conçue pour valider "
    "le pipeline d'entraînement de bout en bout. Voici ses caractéristiques principales :"
)
kv("Espace d'actions",    "MultiDiscrete avec bins continus (throttle ×5, steer ×5, boost ×2, jump ×2, pitch ×5, yaw ×5, roll ×5)")
kv("Observation",         "Vecteur plat de 101 floats (position balle, position voiture, vélocités)")
kv("Fonction de récompense", "Récompenses basiques : vélocité vers la balle, balle vers le but, événement de but")
kv("Adversaire",          "Bot heuristique déterministe (fonce vers la balle)")
kv("Architecture réseau", "MLP simple — couches entièrement connectées")
kv("Framework",           "rlgym-ppo (PPO synchrone, mono-machine)")

heading2("2.2 Limitations identifiées")
body("L'analyse des résultats du jalon 2 a mis en évidence plusieurs problèmes structurels :")
bullet("Espace d'actions redondant : de nombreuses combinaisons de la discrétisation MultiDiscrete sont physiquement invalides ou doublonnées (ex. : boost sans accélération maximale), ce qui gaspille la capacité du réseau.")
bullet("Récompenses trop peu denses : l'agent recevait peu de signal entre les événements de but, entraînant un apprentissage lent et un comportement de rotation circulaire (\"circling\").")
bullet("Adversaire statique : un bot heuristique fixe crée un équilibre d'entraînement artificiellement simple, conduisant à une sur-adaptation à un style de jeu unique.")
bullet("Entropie instable : la politique oscillait sans converger, symptôme d'un signal de récompense insuffisamment informatif.")
bullet("Sous-utilisation du GPU : ~10-20 % d'utilisation, dû au goulot d'étranglement CPU inhérent à la collecte d'expériences en simulation.")

# ══════════════════════════════════════════════════════════════════════════════
#  3. ÉVOLUTIONS V2
# ══════════════════════════════════════════════════════════════════════════════
heading1("3. Architecture version 2 — Améliorations apportées")
separator()

heading2("3.1 Espace d'actions discret pruné (90 actions)")
body(
    "Suite à des recherches dans la communauté rlgym et l'analyse de travaux publiés sur les agents "
    "Rocket League haute performance, nous avons adopté un espace d'actions discret construit "
    "à partir d'une table de correspondance (lookup table) de 90 actions valides, "
    "contre plusieurs centaines dans la version MultiDiscrete."
)
body("La table est construite en deux blocs :")
bullet("Actions au sol (24 actions) : combinaisons de throttle ∈ {-1, 0, 1}, steer ∈ {-1, 0, 1}, boost ∈ {0, 1}, handbrake ∈ {0, 1} — en excluant les combinaisons boost=1 avec throttle≠1 (physiquement inutiles).", level=0)
bullet("Actions aériennes (66 actions) : combinaisons de pitch, yaw, roll ∈ {-1, 0, 1}, jump ∈ {0, 1}, boost ∈ {0, 1} — en excluant jump+yaw≠0 (sideflip non nécessaire) et le cas noop déjà couvert.", level=0)
body(
    "Cette approche réduit la complexité combinatoire, élimine les actions redondantes, "
    "et accélère la convergence de la politique en concentrant la capacité d'expression sur "
    "des actions physiquement significatives. Le réseau de sortie passe de >500 logits à 90."
)

heading2("3.2 Fonction de récompense avancée — 12 composantes")
body(
    "La fonction de récompense a été entièrement redessinée pour fournir un signal dense "
    "et multi-objectif, inspiré des meilleures pratiques observées dans la communauté. "
    "Elle se décompose en deux niveaux : qualité d'état (globale) et qualité individuelle (par joueur)."
)

heading2("  3.2.1 Qualité d'état (state quality)")
bullet("Distance balle–but : récompense différentielle pondérée par exp(-dist/vitesse_max) — plus la balle est proche du but adverse, plus l'état est favorable.")
bullet("Probabilité de victoire : estimée à partir du score courant et du temps restant, intégrée comme signal global de progression de partie.")

heading2("  3.2.2 Récompenses individuelles")
bullet("Distance joueur–balle (Liu distance) : exp(-dist/1410), favorise la proximité sans exiger le contact direct.")
bullet("Alignement joueur–but : cosinus entre le vecteur joueur→balle et joueur→but adverse, encourage le positionnement stratégique.")
bullet("Touch height (hauteur de touche) : activation cubique cbrt((z-150)/CEILING_Z)², récompensant les contacts aériens proportionnellement à la hauteur. Les touches au sol génèrent ~0.04, un dribble soutenu peut atteindre ~20.")
bullet("Facteur de distance au mur : module la récompense de touche selon l'éloignement des murs, favorisant les jeux ouverts.")
bullet("Flip reset : +10 de récompense lorsque l'agent récupère son double-saut en touchant la balle par en-dessous en altitude — encourage une mécanique avancée.")
bullet("Accélération de balle : récompense proportionnelle à ||Δv_balle|| / vitesse_max lors des touches au sol, favorisant les frappes puissantes.")
bullet("Gestion du boost : gain pondéré par sqrt(boost), perte pénalisée proche du sol uniquement — encourage l'économie de boost et la collecte active.")
bullet("Vitesse angulaire : légère récompense pour la rotation en vol, favorisant l'exploration des mécaniques aériennes en début d'entraînement.")
bullet("Pénalité sol : malus si le joueur reste au sol sans activité (touche_grass), décourage la passivité.")
bullet("Démolitions : +demo_w/2 pour le démolisseur, -demo_w/2 pour le démoli — intègre cet aspect tactique du jeu.")

heading2("  3.2.3 Team spirit et punition adversaire")
body(
    "Les récompenses individuelles sont agrégées avec un coefficient de team_spirit = 0.6 : "
    "chaque joueur reçoit 40% de sa récompense individuelle et 60% de la moyenne de son équipe, "
    "encourageant la coopération. Un terme d'opposition (opponent_punish_w = 1.0) soustrait "
    "la récompense moyenne de l'équipe adverse, rendant le jeu explicitement compétitif."
)

heading2("3.3 Self-play par population (Population-Based Self-Play)")
body(
    "L'un des apports les plus significatifs de la v2 est le remplacement de l'adversaire heuristique "
    "par un système de self-play basé sur une population de versions passées de l'agent."
)
body("Fonctionnement du système :")
bullet("L'agent bleu (apprenant) est entraîné par PPO standard.")
bullet("L'agent orange (adversaire) est une version figée (frozen) de l'agent, chargée aléatoirement depuis un pool des 10 derniers checkpoints.")
bullet("L'adversaire est remplacé toutes les 3 000 000 étapes par un checkpoint différent tiré du pool.")
bullet("En l'absence de checkpoint (début d'entraînement), l'agent orange choisit des actions aléatoires.")
body(
    "Cette approche évite la sur-adaptation à un adversaire fixe, expose l'agent à des styles "
    "de jeu variés (y compris ses propres versions passées), et produit une politique plus robuste "
    "et généralisable — tel que démontré empiriquement dans la littérature sur les agents de jeux compétitifs."
)

heading2("3.4 Architecture réseau et hyperparamètres PPO")
body("Le réseau de politique et le critique ont été élargis pour correspondre à la complexité accrue de la tâche :")
kv("Politique (actor)", "DiscreteFF : 101 → 1024 → 1024 → 90 (ReLU activations)")
kv("Critique (critic)",  "DiscreteFF : 101 → 1024 → 1024 → 1 (valeur d'état)")
kv("Taille de batch PPO", "50 000 steps")
kv("Mini-batch",          "10 000 steps")
kv("Epochs PPO",          "3 par itération")
kv("Coefficient d'entropie", "0.01 (exploration maintenue)")
kv("Learning rate",       "5×10⁻⁴ → 5×10⁻⁵ (cosine decay sur 100M steps)")
kv("Standardisation des retours", "Activée (stabilise l'entraînement)")
kv("Steps par itération", "100 000")
kv("Buffer d'expérience", "300 000 steps")
kv("Processus parallèles (N_PROC)", "20 → 40 (après augmentation du fichier de pagination Windows)")
kv("Action repeat",       "8 frames par action")

# ══════════════════════════════════════════════════════════════════════════════
#  4. TABLEAU COMPARATIF
# ══════════════════════════════════════════════════════════════════════════════
heading1("4. Tableau comparatif v1 / v2")
separator()
doc.add_paragraph()

rows = [
    ["Espace d'actions",     "MultiDiscrete binné (~500+ combinaisons)", "Discret pruné — 90 actions valides"],
    ["Observation",          "101 floats (basique)",                     "101 floats (étendu : boost pads, vecteurs relatifs)"],
    ["Récompense",           "3 composantes (vélocité, direction, but)", "12 composantes (dense, multi-objectif)"],
    ["Adversaire",           "Heuristique déterministe",                 "Self-play populationnel (pool 10 checkpoints)"],
    ["Architecture réseau",  "MLP simple",                               "DiscreteFF 101→1024→1024→90/1"],
    ["Team spirit",          "Non implémenté",                           "0.6 (coopération équipe)"],
    ["Flip reset",           "Non récompensé",                           "+10 (mécanique avancée encouragée)"],
    ["Gestion du boost",     "Non récompensée",                          "Gain sqrt-pondéré, perte pénalisée au sol"],
    ["Processus parallèles", "20",                                        "40 (après fix mémoire virtuelle)"],
    ["Learning rate",        "Fixe",                                      "Cosine decay 5e-4 → 5e-5"],
]
add_comparison_table(rows)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  5. RÉSULTATS D'ENTRAÎNEMENT
# ══════════════════════════════════════════════════════════════════════════════
heading1("5. Résultats d'entraînement — Version 2")
separator()

heading2("5.1 Métriques observées")
body(
    "Les métriques ci-dessous sont relevées après environ 8 millions de steps d'entraînement "
    "en v2, soit environ 80 itérations PPO :"
)
kv("Récompense politique",          "1.49 (en progression)")
kv("Entropie",                      "4.36 (stable — exploration maintenue)")
kv("Value Function Loss",           "0.055 (très faible, critique bien calibré)")
kv("KL Divergence moyenne",         "0.0026 (mises à jour conservatrices et stables)")
kv("Steps collectés / seconde",     "~7 500 (collecte) / ~6 200 (global)")
kv("Mises à jour cumulées",         "1 374")

heading2("5.2 Analyse")
body(
    "La Value Function Loss de 0.055 confirme que le critique a rapidement convergé vers "
    "une estimation correcte des retours, ce qui est un indicateur de santé important pour PPO. "
    "L'entropie stable à 4.36 (sur un maximum théorique de ln(90) ≈ 4.50) indique que l'agent "
    "maintient une exploration suffisante sans s'effondrer vers une politique déterministe prématurément."
)
body(
    "La KL divergence très faible (0.0026) confirme que les mises à jour de politique sont "
    "conservatrices et stables, en accord avec la théorie PPO. Aucun collapse d'entropie "
    "ou divergence de valeur n'a été observé."
)
body(
    "Le comportement observé en simulation montre un agent qui commence à se positionner "
    "stratégiquement par rapport à la balle et au but, avec des premiers signes de prise "
    "en compte du positionnement aérien."
)

# ══════════════════════════════════════════════════════════════════════════════
#  6. DÉFIS TECHNIQUES
# ══════════════════════════════════════════════════════════════════════════════
heading1("6. Défis techniques rencontrés")
separator()

heading2("6.1 Goulot d'étranglement CPU / GPU")
body(
    "L'apprentissage par renforcement avec simulation physique (RocketSim) est structurellement "
    "limité par le CPU : la collecte d'expériences est séquentielle par worker et mobilise "
    "le processeur à ~70%, tandis que le GPU n'est sollicité que pendant les phases de mise à jour "
    "PPO (~2.8s sur 16s par itération). Cette alternance produit une utilisation GPU apparente "
    "de 10–30% qui est normale et non pathologique."
)
body(
    "La solution pour augmenter le débit serait un passage à un framework distribué (workers "
    "sur plusieurs machines ou CPU threads natifs en C++), ce qui dépasse le périmètre actuel."
)

heading2("6.2 Mémoire virtuelle Windows (WinError 1455)")
body(
    "Le lancement de 36+ processus Python chargeant chacun les DLLs CUDA (cuBLAS, cuFFT, cuRAND — "
    "environ 2–3 Go de mémoire virtuelle par processus) épuisait le fichier de pagination Windows "
    "par défaut (géré automatiquement à ~36 Go). "
    "Solution appliquée : configuration manuelle du fichier de pagination à 49 152 Mo initial / "
    "65 536 Mo maximum, permettant 40 processus parallèles stables."
)

heading2("6.3 Stabilité initiale du critic")
body(
    "Lors de la première itération d'entraînement v2, une VF Loss de NaN a été observée, "
    "causée par une variance quasi-nulle des retours (tous les états initiaux sont similaires). "
    "La standardisation des retours (standardize_returns=True) a résolu ce problème dès la "
    "deuxième itération, la loss tombant de 3.1 à 2.08 puis 0.177 en trois itérations."
)

# ══════════════════════════════════════════════════════════════════════════════
#  7. CONCLUSION ET PERSPECTIVES
# ══════════════════════════════════════════════════════════════════════════════
heading1("7. Conclusion et perspectives")
separator()

body(
    "La version 2 représente une refonte substantielle du système d'entraînement, "
    "ancrée dans les pratiques de la communauté et validée par les métriques d'entraînement. "
    "Les trois apports majeurs — espace d'actions pruné, récompense dense multi-objectif, "
    "et self-play populationnel — forment un système cohérent dont les effets se cumulent : "
    "l'agent apprend des comportements plus fins grâce au signal de récompense riche, "
    "et est exposé à une plus grande diversité de situations grâce au self-play."
)

body("Perspectives pour la suite du projet :")
bullet("Augmenter le budget de calcul (davantage de steps) pour que le self-play monte en niveau progressivement.")
bullet("Évaluer le niveau atteint contre des bots de référence (All-Star, Psyonix bots) pour situer l'agent sur une échelle de compétence.")
bullet("Explorer un passage à un framework C++ natif (tel que RLGymPPO_CPP) pour un gain de vitesse estimé à ×5–10, rendant possible un entraînement haute performance sur machine personnelle.")
bullet("Affiner les poids de récompense par ablation study une fois l'agent stabilisé.")

doc.add_paragraph()
separator()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("Projet PRL — Rapport Jalon 3  |  Apprentissage par Renforcement appliqué à Rocket League")
set_font(r, size=9, color=LGREY, italic=True)

doc.save("Rapport_Jalon3_PRL.docx")
print("Document généré : Rapport_Jalon3_PRL.docx")
